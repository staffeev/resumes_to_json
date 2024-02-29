import pandas as pd
import os
import fitz 
import docx
import aspose.words as aw


def convert_docx_to_pdf(name, extension):
    doc = aw.Document(f"{name}.{extension}")
    doc.save(f"{name}.pdf")


def remove_tmp_files(filanames):
    [os.remove(name) for name in filanames]


def read(path="source"):
    texts = []
    tmp_files = []
    for filename in map(lambda x: path + "/" + x, os.listdir(path)):
        dot_ix = filename.rindex(".")
        name = filename[:dot_ix]
        extension = filename[dot_ix + 1:]
        if extension in ("doc", "docx"):
            convert_docx_to_pdf(name, extension)
            tmp_files.append(name + ".pdf")
        for block in process_pdf_file(name + ".pdf"):
            texts.append((filename, block))
    remove_tmp_files(tmp_files)
    return pd.DataFrame(texts, columns=["Filename", "Text"])


def process_pdf_file(filename):
    """Получает весь текст из pdf"""
    doc = fitz.open(filename)
    blocks = []
    for page in doc:
        for block in page.get_text_blocks():
            if block[4].startswith("<image:") or "Aspose.Words" in block[4]:
                continue
            blocks.append(block[4])
    return blocks


def process_docx_file(filename):
    """Получает весь текст из docx"""
    doc = docx.Document(filename)
    return "\n".join([i.text for i in doc.paragraphs])


if __name__ == "__main__":
    df = read()
    df.to_csv("csv/text.csv", escapechar="\\")